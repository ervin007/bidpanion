"""
Wrapper module for parsing functionality.
Integrates existing parsing code from parsing.ipynb.
"""

import sys
import os
from pathlib import Path

# Import all parsing functions from the notebook-derived code
import fitz  # PyMuPDF
import pdfplumber
import pytesseract
from PIL import Image
import io
from docx import Document
import openpyxl
import zipfile
import subprocess
import tempfile
import shutil

# RTF support
try:
    from striprtf.striprtf import rtf_to_text
    RTF_AVAILABLE = True
except ImportError:
    RTF_AVAILABLE = False

def extract_text_images_tables(pdf_path, f):
    """Extracts text, images (OCR), and tables from PDF files."""
    doc = fitz.open(pdf_path)
    
    with pdfplumber.open(pdf_path) as pdf:
        f.write(f"\n=== Extracted from {pdf_path} ===\n")

        for page_num, page in enumerate(doc, start=1):
            f.write(f"\n=== Page {page_num} ===\n")

            text = page.get_text("text")
            if text.strip():
                f.write("\n[Text]\n")
                f.write(text + "\n")

            plumber_page = pdf.pages[page_num - 1]
            tables = plumber_page.extract_tables()
            if tables:
                f.write(f"\n=== Tables from Page {page_num} ===\n")
                for table_index, table in enumerate(tables):
                    f.write(f"\n[Table {table_index+1}]\n")
                    for row in table:
                        clean_row = [cell if cell is not None else "" for cell in row]
                        f.write(" | ".join(clean_row) + "\n")

            images = page.get_images(full=True)
            for img_index, img in enumerate(images):
                try:
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_data = base_image["image"]

                    image = Image.open(io.BytesIO(image_data))
                    ocr_text = pytesseract.image_to_string(image, lang="deu")

                    if ocr_text.strip():
                        f.write(f"\n[OCR from Image {img_index+1}]\n")
                        f.write(ocr_text + "\n")
                except Exception as e:
                    print(f"Warning: Could not process image {img_index+1}: {e}")

def extract_text_from_docx(docx_path, f):
    """Extracts text from a Word (.docx) document."""
    doc = Document(docx_path)
    f.write(f"\n=== Extracted from {docx_path} ===\n")

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            f.write(text + "\n")
    
    for table_index, table in enumerate(doc.tables):
        f.write(f"\n[Table {table_index+1}]\n")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            f.write(" | ".join(cells) + "\n")

def extract_text_from_excel(xlsx_path, f):
    """Extracts text from all sheets of an Excel (.xlsx or .xlsm) file."""
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)

    for sheet in wb.sheetnames:
        ws = wb[sheet]
        f.write(f"\n=== Extracted from {xlsx_path} - Sheet: {sheet} ===\n")

        for row in ws.iter_rows(values_only=True):
            clean_row = [str(cell) if cell is not None else "" for cell in row]
            f.write(" | ".join(clean_row) + "\n")

def extract_text_from_rtf(rtf_path, f):
    """Extracts text from RTF files."""
    if not RTF_AVAILABLE:
        print(f"⚠️ Skipping RTF file (striprtf not installed): {rtf_path}")
        return
    
    try:
        with open(rtf_path, 'r', encoding='utf-8', errors='ignore') as rtf_file:
            rtf_content = rtf_file.read()
            text = rtf_to_text(rtf_content)
            f.write(f"\n=== Extracted from {rtf_path} ===\n")
            f.write(text + "\n")
    except Exception as e:
        print(f"Error extracting RTF {rtf_path}: {e}")

def extract_text_from_txt(txt_path, f):
    """Extracts text from plain text files (.txt, .aidf, etc.)."""
    try:
        with open(txt_path, 'r', encoding='utf-8', errors='ignore') as txt_file:
            text = txt_file.read()
            f.write(f"\n=== Extracted from {txt_path} ===\n")
            f.write(text + "\n")
    except Exception as e:
        print(f"Error extracting text from {txt_path}: {e}")

def extract_zip(zip_path, extract_to):
    """Extracts a ZIP file to a temporary directory."""
    try:
        with zipfile.ZipFile(zip_path, 'r') as zip_ref:
            zip_ref.extractall(extract_to)
            print(f"Extracted ZIP: {zip_path} -> {extract_to}")
            return True
    except Exception as e:
        print(f"Error extracting ZIP {zip_path}: {e}")
        return False

def find_7z_executable():
    """Find 7z.exe on Windows. Returns path if found, None otherwise."""
    common_paths = [
        r"C:\Program Files\7-Zip\7z.exe",
        r"C:\Program Files (x86)\7-Zip\7z.exe",
        r"C:\ProgramData\chocolatey\bin\7z.exe",
    ]
    
    for path in common_paths:
        if os.path.exists(path):
            return path
    
    try:
        result = subprocess.run(['where', '7z'], capture_output=True, text=True, shell=True)
        if result.returncode == 0:
            return result.stdout.strip().split('\n')[0]
    except:
        pass
    
    return None

def extract_7z(sevenz_path, extract_to):
    """Extracts a 7z file using system 7z.exe."""
    sevenz_exe = find_7z_executable()
    
    if not sevenz_exe:
        print(f"⚠️ 7-Zip not found. Skipping: {sevenz_path}")
        return False
    
    try:
        result = subprocess.run(
            [sevenz_exe, 'x', sevenz_path, f'-o{extract_to}', '-y'],
            capture_output=True,
            text=True,
            timeout=300
        )
        
        if result.returncode == 0:
            print(f"Extracted 7Z: {sevenz_path} -> {extract_to}")
            return True
        else:
            print(f"Error extracting 7Z {sevenz_path}: {result.stderr}")
            return False
    except Exception as e:
        print(f"Error extracting 7Z {sevenz_path}: {e}")
        return False

def process_file(input_file, f):
    """Process a single file based on its extension."""
    ext = os.path.splitext(input_file)[1].lower()
    
    try:
        if ext == ".pdf":
            extract_text_images_tables(input_file, f)
        elif ext == ".docx":
            extract_text_from_docx(input_file, f)
        elif ext in [".xlsx", ".xlsm"]:
            extract_text_from_excel(input_file, f)
        elif ext == ".rtf":
            extract_text_from_rtf(input_file, f)
        elif ext in [".txt", ".aidf"]:
            extract_text_from_txt(input_file, f)
        else:
            print(f"Skipping unsupported file type: {input_file}")
            return False
        return True
    except Exception as e:
        print(f"Error processing {input_file}: {e}")
        return False

def process_folder_recursively(folder_path, output_file, temp_dirs):
    """Recursively process all files in a folder, including handling ZIP and 7Z files."""
    for root, _, files in os.walk(folder_path):
        for filename in files:
            input_file = os.path.join(root, filename)
            ext = os.path.splitext(filename)[1].lower()
            
            if ext in [".zip", ".7z"]:
                temp_dir = tempfile.mkdtemp()
                temp_dirs.append(temp_dir)
                
                archive_type = "ZIP" if ext == ".zip" else "7Z"
                success = extract_zip(input_file, temp_dir) if ext == ".zip" else extract_7z(input_file, temp_dir)
                
                if success:
                    with open(output_file, "a", encoding="utf-8") as f:
                        f.write(f"\n=== Contents from {archive_type}: {filename} ===\n")
                    process_folder_recursively(temp_dir, output_file, temp_dirs)
            else:
                print(f"Processing {input_file}...")
                with open(output_file, "a", encoding="utf-8") as f:
                    process_file(input_file, f)

def parse_tender_files(input_folder: str, output_folder: str) -> str:
    """
    Main entry point for parsing tender files.
    
    Args:
        input_folder: Path to folder containing raw tender files
        output_folder: Path to output folder for processed files
        
    Returns:
        Path to the merged text file
    """
    input_path = Path(input_folder)
    output_path = Path(output_folder)
    
    if not input_path.exists():
        raise FileNotFoundError(f"Input folder not found: {input_folder}")
    
    output_path.mkdir(parents=True, exist_ok=True)
    
    # Create merged output file
    folder_name = input_path.name if input_path.name != 'raw' else output_path.name
    output_txt = output_path / f"{folder_name}_merged.txt"
    
    temp_dirs = []
    
    try:
        for root, _, files in os.walk(input_folder):
            for filename in files:
                input_file = os.path.join(root, filename)
                ext = os.path.splitext(filename)[1].lower()
                
                if ext in [".zip", ".7z"]:
                    temp_dir = tempfile.mkdtemp()
                    temp_dirs.append(temp_dir)
                    
                    archive_type = "ZIP" if ext == ".zip" else "7Z"
                    print(f"Processing {archive_type}: {input_file}...")
                    
                    success = extract_zip(input_file, temp_dir) if ext == ".zip" else extract_7z(input_file, temp_dir)
                    
                    if success:
                        with open(output_txt, "a", encoding="utf-8") as f:
                            f.write(f"\n=== Contents from {archive_type}: {filename} ===\n")
                        process_folder_recursively(temp_dir, output_txt, temp_dirs)
                else:
                    print(f"Processing {input_file}...")
                    with open(output_txt, "a", encoding="utf-8") as f:
                        process_file(input_file, f)
        
        print(f"✅ Extraction complete! Saved to: {output_txt}")
        return str(output_txt)
    
    finally:
        for temp_dir in temp_dirs:
            try:
                shutil.rmtree(temp_dir)
                print(f"Cleaned up temp directory: {temp_dir}")
            except Exception as e:
                print(f"Error cleaning up {temp_dir}: {e}")
